import streamlit as st
import random
import asyncio
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
from langchain.prompts import PromptTemplate
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain.chains import RetrievalQA
from langchain.schema import Document
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Define the custom prompt template for the RetrievalQA chain
custom_prompt_template = """
Use the following pieces of information to answer the user's question in a comprehensive and informative way, leveraging the retrieved documents. If the retrieved documents are not helpful, try to provide relevant information or rephrase the question.

Context: {context}
Question: {question}

Only return the helpful answer below and nothing else.
Helpful answer:
"""

# Initialize the PromptTemplate with the custom prompt
prompt = PromptTemplate(template=custom_prompt_template, input_variables=['context', 'question'])

# Function to create a RetrievalQA chain
def retrieval_qa_chain(llm, db):
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type='stuff',
        retriever=db.as_retriever(search_kwargs={'k': 5}),  # Retrieve top 5 documents
        return_source_documents=False,  # Do not return source documents
        chain_type_kwargs={'prompt': prompt}
    )
    return qa_chain

# Function to load the language model (LLM)
def load_llm():
    llm = ChatGroq(
        groq_api_key="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",
        model_name="Llama3-8b-8192",
        max_tokens=8192,
        temperature=0.5
    )
    return llm

# Function to create vector database from in-memory file
def create_vector_db_from_memory(file_bytes, file_type):
    # Convert file bytes to text based on file type
    if file_type == 'pdf':
        pdf = PdfReader(BytesIO(file_bytes))
        text = ""
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            text += page.extract_text()
    elif file_type == 'docx':
        doc = DocxDocument(BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    elif file_type == 'txt':
        text = file_bytes.decode('utf-8')
    else:
        raise ValueError("Unsupported file type")

    # Create Document objects
    documents = [Document(page_content=text)]

    # Initialize embeddings
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={'device': 'cpu'})

    # Create FAISS index
    vectorstore = FAISS.from_documents(documents, embeddings)
    return vectorstore
responses = [
    "Here's the information I found for you. If you need anything else, just let me know!",
    "I've got some details on that topic. Feel free to ask if you'd like more information or have other questions!",
    "I hope this answers your question. If you'd like more details or have any other inquiries, I'm here to help!",
    "Here are the details I have for you. If you have more questions or need further assistance, just ask!",
    "I've gathered the relevant information for you. Let me know if there's anything more you'd like to know or if you have any other questions!"
]
# Asynchronous function to initialize the QA bot
async def qa_bot():
    llm = load_llm()
    return llm

# Function to get the QA chain
def get_chain(vectorstore):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    llm = loop.run_until_complete(qa_bot())
    return retrieval_qa_chain(llm, vectorstore)

# Function to handle the QA chain invocation and context-based queries
def new_func(chain, question_text, context=None):
    try:
        if context:
            question_text = f"{context}\n{question_text}"
        result = chain.invoke(question_text)
        return result
    except Exception as e:
        return f"Oops! Something went wrong: {str(e)}"

# Custom CSS for dark mode and styled components
st.markdown("""
    <style>
    .main {
        background-color: #121212;
        color: #ffffff;
        font-family: Arial, sans-serif;
    }
    .header {
        display: flex;
        align-items: center;
        margin-bottom: 20px;
    }
    .header img {
        height: 60px; /* Adjust height as needed */
        margin-right: 20px;
    }
    .header h1 {
        font-size: 3em;
        color: #4187CE;
        margin: 0;
    }
    .input-section {
        margin: 20px 0;
        padding: 20px;
        background-color: #2C2C2C;
        border-radius: 8px;
        box-shadow: 0px 4px 8px rgba(0, 0, 0, 0.1);
    }
    .history div {
        margin-bottom: 10px;
        display: flex;
        align-items: center;
    }
    .user {
        font-weight: bold;
        color: white;
        justify-content: flex-end;
    }
    .bot {
        font-weight: bold;
        color: white;
        justify-content: flex-start;
    }
    .message {
        max-width: 70%;
        padding: 10px;
        border-radius: 8px;
    }
    .user .message {
        background-color: #BD1362;
        color: #ffffff;
        align-self: flex-end;
    }
    .bot .message {
        background-color: #5FB233;
        color: #ffffff;
        align-self: flex-start.
    }
    </style>
    """, unsafe_allow_html=True)

# Header section with logo and title
st.image("./images/godrej.png", width=100)  # Adjust width as needed
st.markdown("""
    <div class="header">
        <h1>Godrej Guide Support Chat</h1>
    </div>
    """, unsafe_allow_html=True)

# Greet the user
st.write('Hello there! 👋 Welcome to the support chat. How can I assist you today?')

# Initialize session state for chat history and context
if 'history' not in st.session_state:
    st.session_state['history'] = []
    st.session_state['context'] = ""

# Create a form for user input and submission
with st.form(key='chat_form', clear_on_submit=True):
    # Text input for user's question
    user_input = st.text_input("Your Question:", key="input_field")
    # Submit button
    submit_button = st.form_submit_button(label='Submit')

    if submit_button and user_input:
        # Get the vector store from uploaded file
        vectorstore = st.session_state.get('vectorstore', None)
        
        if vectorstore is None:
            st.write("Please upload a file to create the vector store.")
        else:
            # Get the QA chain
            chain = get_chain(vectorstore)
            # Check if elaboration is needed based on user input
            elaboration_needed = "please elaborate" in user_input.lower() or "expand on that" in user_input.lower() or "give more details" in user_input.lower()
            
            if elaboration_needed:
                # If elaboration is needed, use the last bot response as context
                if len(st.session_state['history']) > 0 and 'bot' in st.session_state['history'][-1]:
                    last_bot_response = st.session_state['history'][-1]['bot']
                    question_text = f"Can you provide more details or context on this: {last_bot_response}?"
                    st.session_state['context'] += f"\nElaboration on: {last_bot_response}"
                else:
                    question_text = "I don't have a previous response to elaborate on. Could you please ask a different question?"
            else:
                # Otherwise, use the current user input
                question_text = user_input

            # Add user input to chat history
            st.session_state['history'].append({"user": user_input})

        try:
    # Get the response from the QA chain
             res = new_func(chain, question_text, st.session_state['context'])
    
    # Determine how to extract the answer
             if isinstance(res, str):
               answer = res
             elif isinstance(res, dict) and "result" in res:
               answer = res.get("result", "No result found")
             else:
               answer = "No result found"
        except Exception as e:
               st.write("Oops! Something went wrong:", str(e))
               answer = "Sorry, there was an issue with your request. Please try again later."

# Add bot response to chat history and append a friendly message
st.session_state['history'].append({"bot": answer})
answer += "\n" + random.choice(responses)
st.write(answer)
# Display chat history
if 'history' in st.session_state:
    st.write("<div class='history'><h3>Chat History:</h3>", unsafe_allow_html=True)
    for entry in st.session_state['history']:
        if 'user' in entry:
            st.write(f"<div class='user'><div class='user'>User</div><div class='message'>{entry['user']}</div></div>", unsafe_allow_html=True)
            st.write(f"<br>", unsafe_allow_html=True)
        if 'bot' in entry:
            st.write(f"<div class='bot'><div class='bot'>Godrej Guide</div><div class='message'>{entry['bot']}</div></div>", unsafe_allow_html=True)
            st.write(f"<br>", unsafe_allow_html=True)
    st.write("</div>", unsafe_allow_html=True)

st.write("## Upload File")
uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "txt"])

if uploaded_file is not None:
    st.write("File uploaded:", uploaded_file.name)
    
    # Get file bytes and type
    file_bytes = uploaded_file.read()
    file_type = uploaded_file.type.split('/')[1]
    
    # Create vector database from uploaded file
    vectorstore = create_vector_db_from_memory(file_bytes, file_type)
    st.session_state['vectorstore'] = vectorstore
    st.write("Vector database created from the uploaded file.")
